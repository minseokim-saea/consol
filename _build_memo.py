"""IT부서 설명자료 Word 문서 생성기."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

DOC_PATH = r'C:\패키지프로그램\연결재무보고시스템_IT부서_설명자료.docx'

FONT_KR = '맑은 고딕'    # Malgun Gothic — Windows 기본 한글 글꼴
FONT_EN = 'Calibri'

doc = Document()

# ── 페이지 설정 ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# ── 기본 스타일 ──────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = FONT_KR
style_normal.font.size = Pt(10.5)
# 한글 폰트 명시 (EastAsia)
rPr = style_normal.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), FONT_KR)
rFonts.set(qn('w:ascii'), FONT_EN)
rFonts.set(qn('w:hAnsi'), FONT_EN)


def _apply_font(run, size=None, bold=False, color=None, font_kr=FONT_KR, font_en=FONT_EN):
    """run에 한글/영문 폰트와 속성 적용."""
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_kr)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    if size is not None:
        run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color is not None:
        run.font.color.rgb = color


def add_heading(text, level=1, color=None):
    """레벨별 헤딩 추가. level 0=타이틀, 1=대분류, 2=중분류."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    if level == 0:
        run = p.add_run(text)
        _apply_font(run, size=18, bold=True,
                    color=color or RGBColor(0x1F, 0x38, 0x64))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run = p.add_run(text)
        _apply_font(run, size=14, bold=True,
                    color=color or RGBColor(0x1F, 0x38, 0x64))
        # 하단 보더
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '2')
        bottom.set(qn('w:color'), '4472C4')
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        run = p.add_run(text)
        _apply_font(run, size=11.5, bold=True,
                    color=color or RGBColor(0x2F, 0x55, 0x97))
    return p


def add_paragraph(text, bold=False, size=10.5, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    _apply_font(run, size=size, bold=bold)
    return p


def add_bullet(text, level=0):
    """글머리기호 항목."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    # remove default run, add ours
    while p.runs:
        p.runs[0].text = ''
        break
    run = p.add_run(text)
    _apply_font(run, size=10.5)
    return p


def _set_cell_border(cell, **kwargs):
    """셀 테두리 설정. kwargs: top/bottom/left/right 각각 {'sz':, 'color':, 'val':}."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        opts = kwargs.get(edge) or {'sz': 4, 'val': 'single', 'color': 'BFBFBF'}
        existing = tcBorders.find(qn(f'w:{edge}'))
        if existing is not None:
            tcBorders.remove(existing)
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), opts.get('val', 'single'))
        b.set(qn('w:sz'), str(opts.get('sz', 4)))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), opts.get('color', 'BFBFBF'))
        tcBorders.append(b)


def _set_cell_shading(cell, fill):
    """셀 배경색."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _style_cell(cell, text, *, bold=False, size=10.5, align=None, header=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _apply_font(run, size=size, bold=bold or header,
                color=color or (RGBColor(0xFF, 0xFF, 0xFF) if header else None))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_border(cell)
    if header:
        _set_cell_shading(cell, '1F3864')
    return cell


def add_table(headers, rows, col_widths_cm=None):
    """헤더가 있는 2열 이상 표 추가."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths_cm:
        # 표 전체와 각 컬럼 너비
        for ci, w in enumerate(col_widths_cm):
            for r in range(len(rows) + 1):
                table.cell(r, ci).width = Cm(w)

    # 헤더 행
    for ci, h in enumerate(headers):
        _style_cell(table.cell(0, ci), h, header=True, align='center')
    # 데이터 행
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            _style_cell(table.cell(ri, ci), val)
            # 줄무늬
            if ri % 2 == 0:
                _set_cell_shading(table.cell(ri, ci), 'F2F2F2')

    # 표 후 간격
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ════════════════════════════════════════════════════════════
# 문서 본문 작성
# ════════════════════════════════════════════════════════════

# ── 헤더 / 타이틀 ───────────────────────────────────────────
add_heading('연결재무보고 통합 시스템', level=0)
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(8)
sub_run = sub_p.add_run('— 사내 서버 배포를 위한 IT 부서 안내자료 —')
_apply_font(sub_run, size=11.5, color=RGBColor(0x6C, 0x75, 0x7D))

# 메타 정보 표 (제목 아래)
meta = doc.add_table(rows=2, cols=4)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
meta_widths = [3.0, 5.0, 3.0, 5.0]
for ci, w in enumerate(meta_widths):
    meta.cell(0, ci).width = Cm(w)
    meta.cell(1, ci).width = Cm(w)
_style_cell(meta.cell(0, 0), '문서 제목', header=True, align='center')
_style_cell(meta.cell(0, 1), '연결재무보고 통합 시스템 도입 안내', align='left', bold=True)
_style_cell(meta.cell(0, 2), '작성일', header=True, align='center')
_style_cell(meta.cell(0, 3), '2026. 5. 14.', align='left')
_style_cell(meta.cell(1, 0), '작성 부서', header=True, align='center')
_style_cell(meta.cell(1, 1), '재무팀', align='left')
_style_cell(meta.cell(1, 2), '수신', header=True, align='center')
_style_cell(meta.cell(1, 3), 'IT 부서', align='left')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── 1. 문서 목적 ────────────────────────────────────────────
add_heading('1. 문서 목적', level=1)
add_paragraph(
    '본 자료는 재무팀에서 자체 개발한 「연결재무보고 통합 시스템」(이하 "본 시스템")의 '
    '도입 배경과 주요 기능, 기대효과를 정리하여, 사내 운영 서버 배포에 필요한 IT 부서의 '
    '협조를 요청드리기 위해 작성하였습니다.'
)
add_paragraph(
    '본 시스템은 분기마다 반복되는 다수 자회사의 연결결산 업무를 자동화·표준화하기 위한 '
    '사내 웹 애플리케이션으로, 외부 서비스 의존 없이 사내망에서 단독 운영이 가능합니다.'
)

# ── 2. 시스템 개요 ──────────────────────────────────────────
add_heading('2. 시스템 개요', level=1)
add_paragraph(
    '본 시스템은 다음의 결산 절차를 하나의 웹 화면에서 일원화하여 처리합니다.'
)
add_bullet('자회사 결산패키지(엑셀) 업로드 및 기간별 보관')
add_bullet('패키지 데이터 자동 합산 (BS / PL / 현금흐름표)')
add_bullet('연결조정 분개 입력 및 자동 검증 (차변·대변 균형, 코드 누락 차단)')
add_bullet('포함 그룹(다단계 연결)의 재귀 합산 처리')
add_bullet('최종 연결재무제표 엑셀 자동 생성')
add_bullet('그룹별 핵심지표 대시보드 (KPI · 도넛 · 막대 차트)')
add_paragraph(
    '분기별 결산 사이클에 맞춰 활용하며, 다수 자회사를 보유한 그룹사의 연결결산 업무에 '
    '최적화되어 있습니다.'
)

# ── 3. 도입 필요성 (강조) ───────────────────────────────────
add_heading('3. 도입 필요성', level=1)

add_heading('3.1 현행 업무의 비효율', level=2)
add_bullet('30개 이상의 자회사 결산패키지를 매 분기마다 엑셀로 수기 합산하고 있어 단순 반복 작업에 다수의 인시(人時)가 소요됩니다.')
add_bullet('합산 이후 연결조정 분개를 별도 파일에서 수작업으로 입력하며, 차변·대변 균형 확인도 수기로 진행되어 검증 누락의 위험이 상존합니다.')
add_bullet('자회사별 패키지 양식 차이 및 계정코드 불일치로 인한 데이터 정합성 검토에 적지 않은 시간이 투입됩니다.')
add_bullet('결산 결과는 보고용 자료(대시보드·추이·그룹별 KPI 등)로 다시 가공하는 별도 작업이 필요합니다.')
add_bullet('그룹사 계층 구조(예: 글로벌세아 ⊃ 상역 ⊃ 태림 / GIT)로 인한 다단계 연결을 수작업으로 끌어와 합산하는 과정에서 누락 또는 중복 발생 위험이 있습니다.')

add_heading('3.2 업무 리스크', level=2)
add_bullet('수작업 합산에 따른 휴먼 에러로 결산 신뢰성 저하 가능')
add_bullet('분기마다 담당자별 양식·계산식 편차 발생 → 추세 분석의 일관성 훼손')
add_bullet('외부감사·내부감사 대응 시 분개·조정 이력의 추적이 곤란')
add_bullet('결산 일정 압박 시 검증 단계가 단축되어 오류 발견 시점이 늦어짐')
add_bullet('담당자 변경·인수인계 시 학습 곡선이 가파르고, 노하우의 속인화(屬人化)가 큰 편')

# ── 4. 시스템 주요 기능 ────────────────────────────────────
add_heading('4. 시스템 주요 기능', level=1)
add_paragraph('본 시스템이 제공하는 주요 기능은 아래와 같습니다.')
add_table(
    headers=['기능 영역', '핵심 기능'],
    rows=[
        ['패키지 관리',
         '회사별·기간별 엑셀 패키지 업로드, 중복 검출, 매칭 회사 자동 표시, 미업로드 회사 알림'],
        ['자동 합산',
         'BS·PL·현금흐름표 시트별 회사 간 합산, WCE 자본조정 자동 반영, 환산 대체분개 처리'],
        ['연결조정 분개',
         '엑셀 양식 업로드 또는 화면 수기 입력, 차변·대변 합계 자동 검증, 금액 입력 후 코드 누락 셀 업로드 거부, 원본 분개 파일 재다운로드 기능'],
        ['자동 매듭',
         '연결조정·내부거래 분개 적용 시 BS-PL 간 잔액 불일치를 자동 매듭 분개로 보정하여 시산 균형 확보'],
        ['그룹 계층 처리',
         '포함 그룹(rollup)의 결과를 상위 그룹의 한 컬럼으로 자동 끌어오며 다단계 재귀 처리'],
        ['최종 산출물',
         '분개 적용 후 연결재무제표 엑셀 파일 자동 생성·다운로드, 시트별 합산·조정·최종값 컬럼 분리'],
        ['그룹별 대시보드',
         '자산·부채·자본·매출·영업이익·당기순이익·총차입금·은행차입금 KPI 카드, 회사별 자산/매출 구성 도넛 차트, 자산·부채·자본 구성 막대 차트, 합산→최종 비교 막대 차트'],
        ['전년 비교 (YoY)',
         '핵심지표의 전년 연결값을 수기 입력하면 자동 YoY 비교 (BS는 전년말 1세트, PL은 분기별 YTD 누계 4세트)'],
        ['사용자 관리',
         '로그인, 비밀번호 해시 저장, 관리자·일반 사용자 권한 분리, 비밀번호 변경 화면'],
    ],
    col_widths_cm=[4.5, 12.5],
)

# ── 5. 기대효과 (강조) ─────────────────────────────────────
add_heading('5. 기대효과', level=1)

add_heading('5.1 정량적 기대효과', level=2)
add_bullet('결산 소요시간의 의미 있는 단축 — 합산·연결조정·검증·자료작성 단계가 한 화면에서 일원화되어, 분기당 수일 단위의 업무 시간 절감이 기대됩니다.')
add_bullet('합산 단계의 휴먼 에러 발생률 감소 — 동일 양식·동일 계산식이 시스템에 의해 강제 적용됩니다.')
add_bullet('분개 검증의 자동화 — 차변·대변 불균형 및 코드 누락이 업로드 시점에 자동 차단되어 사후 정정 비용이 절감됩니다.')
add_bullet('보고자료 작성 부담 경감 — 대시보드가 결산 직후 자동 생성되어 별도 보고서 가공 시간을 줄입니다.')

add_heading('5.2 정성적 기대효과', level=2)
add_bullet('데이터 일관성·신뢰성 확보 — 분기 간, 그룹 간 동일 기준으로 산출되어 추세 분석과 비교 분석의 신뢰도가 향상됩니다.')
add_bullet('경영진 의사결정 적시 지원 — 결산 결과가 즉시 시각화되어 그룹별 KPI를 실시간 확인할 수 있습니다.')
add_bullet('내·외부 감사 대응력 강화 — 분개 변경 이력 및 자동 매듭 이력이 보존되어 추적성이 높아집니다.')
add_bullet('업무의 표준화 — 담당자가 바뀌어도 동일한 UI·동일한 절차로 운영되어 인수인계 부담이 감소합니다.')
add_bullet('속인화 해소 — 합산·조정 로직이 시스템에 코드화되어 개인 노하우 의존도가 낮아집니다.')
add_bullet('전년 동기 대비 YoY 분석의 자동화 — 입력해 둔 전년 핵심지표와 자동 비교되어 추이 모니터링이 손쉬워집니다.')

# ── 6. 기술 개요 ───────────────────────────────────────────
add_heading('6. 기술 개요', level=1)
add_paragraph('본 시스템은 일반적이고 검증된 오픈소스 기술 스택을 사용하여 사내 운영에 부담이 적습니다.')
add_table(
    headers=['구분', '내용'],
    rows=[
        ['백엔드', 'Python 3.11 + Flask 웹 프레임워크 (경량, 단일 프로세스 운영 가능)'],
        ['데이터 영속', 'JSON 파일 (그룹 정의, 분개, 전년값, 사용자) + 엑셀 파일 (업로드/결과)'],
        ['클라이언트', 'Bootstrap 5 + Chart.js — 별도 설치 불필요, 일반 브라우저에서 동작'],
        ['인증', '세션 기반 로그인, 비밀번호 해시 저장 (단방향)'],
        ['파일 처리', 'openpyxl · pandas — 엑셀 읽기/쓰기'],
        ['외부 의존', '외부 API 호출 없음. 사내 폐쇄망에서도 운영 가능'],
        ['운영 부담', '소수 사용자(재무팀 내) · 결산기 집중 사용 — 단일 머신으로 충분'],
    ],
    col_widths_cm=[4.5, 12.5],
)

# ── 7. 운영 서버 요청 사항 ─────────────────────────────────
add_heading('7. 운영 서버 요청 사항', level=1)
add_paragraph('운영 환경 배포에 필요한 사양 및 권장 구성은 다음과 같습니다.')
add_table(
    headers=['항목', '요구사항 / 권장사항'],
    rows=[
        ['운영체제(OS)', 'Linux (Ubuntu 22.04 LTS 이상) 또는 Windows Server'],
        ['Python', '3.11 이상'],
        ['메모리', '최소 2 GB (4 GB 권장)'],
        ['디스크', '50 GB 이상 (업로드·결과 파일 누적 대비)'],
        ['네트워크', '사내망 한정, 외부 인터넷 차단, 사내 HTTPS 인증서 사용'],
        ['권장 구성', 'gunicorn(WSGI) + nginx(리버스 프록시) 또는 동등 구성'],
        ['백업', 'uploads/ , results/ , *.json 디렉토리에 대한 일 단위 자동 백업'],
        ['로그', '표준 웹 액세스 로그 및 시스템 자체 분개 변경 이력 보관'],
    ],
    col_widths_cm=[4.5, 12.5],
)

# ── 8. 보안·운영 고려사항 ──────────────────────────────────
add_heading('8. 보안·운영 고려사항', level=1)
add_bullet('접근 통제 — 사내 IP 화이트리스트 또는 SSO 연동 검토 가능합니다.')
add_bullet('비밀번호 정책 — 관리자 페이지에서 강제 변경 및 권한 부여가 가능하며, 비밀번호는 해시로만 저장됩니다.')
add_bullet('업로드 제한 — 파일당 50 MB, 확장자(.xlsx / .xlsm)만 허용합니다.')
add_bullet('세션 키 관리 — 무작위로 생성되어 디스크에 저장되며, 서버 재시작 후에도 세션이 유지됩니다.')
add_bullet('외부 통신 없음 — 결산 데이터가 외부로 송출되지 않으며, 폐쇄망 운영이 가능합니다.')
add_bullet('변경 이력 — 분개·전년값 등 핵심 데이터의 수정 시간과 작업자가 자동 기록됩니다.')

# ── 9. 도입 일정 ───────────────────────────────────────────
add_heading('9. 도입 일정 (제안)', level=1)
add_table(
    headers=['단계', '주요 과업', '예상 소요'],
    rows=[
        ['1단계', '사전 협의 — IT 부서와 인프라·보안 사양 확정', '1 ~ 2주'],
        ['2단계', '스테이징 환경 배포 — 시범 사용자 계정 발급, 기능 검증', '1주'],
        ['3단계', '운영 환경 배포 — 사내망 게시, 기존 결산자료 이관', '1주'],
        ['4단계', '정기 결산 시점부터 본격 사용 개시', '결산기 도래 시'],
    ],
    col_widths_cm=[2.5, 11.0, 3.5],
)

# ── 10. 협조 요청 ──────────────────────────────────────────
add_heading('10. IT 부서 협조 요청 사항', level=1)
add_paragraph('아래 사항에 대한 검토와 지원을 요청드립니다.')
add_bullet('운영 서버 1대 할당 (위 7번 사양 기준)')
add_bullet('사내 HTTPS 인증서 발급 및 적용')
add_bullet('사내망 접근 규칙 설정 (외부 차단, 재무팀 및 승인 사용자만 접근)')
add_bullet('백업 정책 수립 협조 (uploads / results / JSON 디렉토리 대상)')
add_bullet('필요 시 사내 도메인 부여 (예: consol.company.local)')
add_bullet('초기 배포 시 가상환경 구성 및 서비스 등록 협조 (재무팀 측 인수 지원 가능)')

# ── Footer 영역 ────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(12)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(20)
run = footer_p.add_run('— 이상 —')
_apply_font(run, size=10, color=RGBColor(0x6C, 0x75, 0x7D))

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_before = Pt(4)
run = contact_p.add_run('문의 : 재무팀  ·  작성일 2026. 5. 14.')
_apply_font(run, size=9.5, color=RGBColor(0x6C, 0x75, 0x7D))

# 저장
doc.save(DOC_PATH)
print(f'OK: {DOC_PATH}')
