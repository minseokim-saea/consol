# -*- coding: utf-8 -*-
"""Google OTP(2단계 인증) 등록 안내 워드 문서 생성 (②⑥ 화면 일러스트 삽입)"""
import os, math
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===================================================================
#  PART 1.  화면 일러스트 PNG 렌더링 (Pillow)
# ===================================================================
S = 2  # 해상도 배율
NAVY1 = (0x1F, 0x38, 0x64)
NAVY2 = (0x2D, 0x5A, 0xA0)
BG = (0xF0, 0xF4, 0xF8)
WHITE = (255, 255, 255)
MUTED = (0x6C, 0x75, 0x7D)
DARKT = (0x21, 0x25, 0x29)
BORDER = (0xDE, 0xE2, 0xE6)
IMGDIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_otp_img")
os.makedirs(IMGDIR, exist_ok=True)

FREG = r"C:\Windows\Fonts\malgun.ttf"
FBLD = r"C:\Windows\Fonts\malgunbd.ttf"
FMONO = r"C:\Windows\Fonts\consola.ttf"

def font(px, bold=False, mono=False):
    path = FMONO if mono else (FBLD if bold else FREG)
    return ImageFont.truetype(path, int(px))

def diag_gradient(w, h, c1, c2):
    """135deg 대각선 그라데이션 (작게 그려 확대)"""
    small = Image.new("RGB", (64, 64))
    px = small.load()
    for y in range(64):
        for x in range(64):
            t = (x + y) / 126.0
            px[x, y] = tuple(int(c1[i] + (c2[i] - c1[i]) * t) for i in range(3))
    return small.resize((w, h), Image.BILINEAR)

def rounded_mask(w, h, r):
    m = Image.new("L", (w, h), 0)
    d = ImageDraw.Draw(m)
    d.rounded_rectangle([0, 0, w - 1, h - 1], radius=r, fill=255)
    return m

def text_center(d, cx, y, s, fnt, fill):
    bb = d.textbbox((0, 0), s, font=fnt)
    d.text((cx - (bb[2] - bb[0]) / 2, y), s, font=fnt, fill=fill)
    return bb[3] - bb[1]

def button_label(d, cx, cy, text, fnt, color=WHITE):
    """체크(✓ 직접 그림) + 텍스트를 가운데 정렬로 버튼에 그림"""
    bb = d.textbbox((0, 0), text, font=fnt)
    tw, th = bb[2] - bb[0], bb[3] - bb[1]
    chk = int(th * 1.0); gap = int(th * 0.45)
    total = chk + gap + tw
    x0 = cx - total / 2
    # 체크마크 (꺾은 선 2개)
    cy0 = cy + th * 0.5
    lw = max(2, int(th * 0.16))
    d.line([(x0, cy0 + chk*0.12), (x0 + chk*0.38, cy0 + chk*0.5)], fill=color, width=lw)
    d.line([(x0 + chk*0.38, cy0 + chk*0.5), (x0 + chk*0.95, cy0 - chk*0.45)], fill=color, width=lw)
    d.text((x0 + chk + gap, cy), text, font=fnt, fill=color)

def lang_pill(card, x, y):
    """우상단 언어 전환 알약 (한국어 | English)"""
    d = ImageDraw.Draw(card)
    f = font(11 * S)
    pad = 6 * S
    t1, t2 = "한국어", "English"
    w1 = d.textbbox((0, 0), t1, font=f)[2]
    w2 = d.textbbox((0, 0), t2, font=f)[2]
    h = 18 * S
    seg1 = pad + w1 + pad
    total = seg1 + pad + w2 + pad
    rx = x - total
    ty = y + (h - f.getbbox(t1)[3]) / 2 - 1 * S
    # 알약 외곽선 + 활성(한국어) 흰 배경
    d.rounded_rectangle([rx, y, rx + total, y + h], radius=h // 2, outline=WHITE, width=max(1, S))
    d.rounded_rectangle([rx, y, rx + seg1, y + h], radius=h // 2, fill=WHITE)
    d.text((rx + pad, ty), t1, font=f, fill=NAVY1)        # 한국어(활성): 남색
    d.text((rx + seg1 + pad, ty), t2, font=f, fill=WHITE)  # English: 흰색

def shield_lock(card, cx, cy, size, color=WHITE, lockcol=NAVY1):
    """방패+자물쇠 아이콘"""
    d = ImageDraw.Draw(card)
    w = size; h = size * 1.15
    top = cy - h / 2
    pts = [(cx - w/2, top + h*0.06), (cx, top), (cx + w/2, top + h*0.06),
           (cx + w/2, top + h*0.5), (cx, top + h), (cx - w/2, top + h*0.5)]
    d.polygon(pts, fill=color)
    # 자물쇠
    bw = w * 0.34; bh = h * 0.26
    bx0 = cx - bw/2; by0 = top + h*0.42
    d.rounded_rectangle([bx0, by0, bx0 + bw, by0 + bh], radius=bw*0.12, fill=lockcol)
    # 고리(shackle)
    sr = bw * 0.32
    d.arc([cx - sr, by0 - sr*1.1, cx + sr, by0 + sr*0.7], 180, 360,
          fill=lockcol, width=max(2, int(size*0.045)))
    # 열쇠구멍
    d.ellipse([cx - bw*0.09, by0 + bh*0.3, cx + bw*0.09, by0 + bh*0.5], fill=color)

def fake_qr(side, modules=25):
    """진짜처럼 보이는 QR 패턴 (파인더 3개 + 의사난수 모듈)"""
    img = Image.new("RGB", (side, side), WHITE)
    d = ImageDraw.Draw(img)
    m = side / modules
    def filled(r, c):
        # 파인더 영역 제외, 결정적 의사난수
        v = (r * 73856093) ^ (c * 19349663) ^ ((r + c) * 83492791)
        return (v >> 5) & 1
    for r in range(modules):
        for c in range(modules):
            in_finder = ((r < 7 and c < 7) or (r < 7 and c >= modules-7)
                         or (r >= modules-7 and c < 7))
            if in_finder:
                continue
            if 8 <= r <= 9 or 8 <= c <= 9:  # 약간의 타이밍 느낌
                pass
            if filled(r, c):
                d.rectangle([c*m, r*m, (c+1)*m, (r+1)*m], fill=(0, 0, 0))
    def finder(r0, c0):
        x0, y0 = c0*m, r0*m
        d.rectangle([x0, y0, x0+7*m, y0+7*m], fill=(0, 0, 0))
        d.rectangle([x0+m, y0+m, x0+6*m, y0+6*m], fill=WHITE)
        d.rectangle([x0+2*m, y0+2*m, x0+5*m, y0+5*m], fill=(0, 0, 0))
    finder(0, 0); finder(0, modules-7); finder(modules-7, 0)
    return img

def dashed_rect(d, box, color, width=1, dash=6, gap=4):
    x0, y0, x1, y1 = box
    def hline(y):
        x = x0
        while x < x1:
            d.line([x, y, min(x+dash, x1), y], fill=color, width=width)
            x += dash + gap
    def vline(x):
        y = y0
        while y < y1:
            d.line([x, y, x, min(y+dash, y1)], fill=color, width=width)
            y += dash + gap
    hline(y0); hline(y1); vline(x0); vline(x1)

def make_card(card_w, body_render, header_render, header_h):
    """카드(헤더+본문)를 BG 위에 얹어 스크린샷처럼 생성. body_render는 높이 반환."""
    pad = 28 * S
    # 우선 본문 높이 측정용 임시
    tmp = Image.new("RGB", (card_w, 4000), WHITE)
    body_h = body_render(tmp, measure=True)
    card_h = header_h + body_h
    card = Image.new("RGB", (card_w, card_h), WHITE)
    # 헤더
    header = diag_gradient(card_w, header_h, NAVY1, NAVY2)
    card.paste(header, (0, 0))
    header_render(card)
    # 본문
    body_render(card, measure=False)
    # 라운드 처리
    r = 16 * S
    mask = rounded_mask(card_w, card_h, r)
    rounded = Image.new("RGB", (card_w, card_h), BG)
    rounded.paste(card, (0, 0), mask)
    # 캔버스(배경 + 그림자 느낌)
    canvas = Image.new("RGB", (card_w + pad*2, card_h + pad*2), BG)
    # 살짝 그림자
    shadow = Image.new("RGBA", canvas.size, (0, 0, 0, 0))
    sd = ImageDraw.Draw(shadow)
    sd.rounded_rectangle([pad+4*S, pad+8*S, pad+card_w+4*S, pad+card_h+8*S],
                         radius=r, fill=(31, 56, 100, 40))
    canvas.paste(Image.alpha_composite(canvas.convert("RGBA"), shadow).convert("RGB"), (0, 0))
    canvas.paste(rounded, (pad, pad), mask)
    return canvas

# ---------- 화면 ② : 2단계 인증(OTP) 등록 ----------
def render_setup():
    card_w = 470 * S
    header_h = 70 * S
    def header(card):
        d = ImageDraw.Draw(card)
        shield_lock(card, 30*S, 35*S, 26*S)
        d.text((52*S, 24*S), "2단계 인증(OTP) 등록", font=font(16*S, bold=True), fill=WHITE)
        lang_pill(card, card_w - 12*S, 12*S)
    def body(card, measure=False):
        pad = 30 * S
        x = pad
        y = header_h + 24*S if not measure else 24*S
        d = ImageDraw.Draw(card)
        if measure:
            y = 0
        else:
            y = header_h + 22*S
        cur = y
        # 호환 안내
        if not measure:
            d.text((x, cur), "Google Authenticator · Microsoft Authenticator · Authy 등과 호환됩니다.",
                   font=font(10.5*S), fill=MUTED)
        cur += 24*S
        # ① 안내
        if not measure:
            d.text((x, cur), "① 인증 앱으로 아래 QR 코드를 스캔하세요",
                   font=font(12*S, bold=True), fill=DARKT)
        cur += 28*S
        # QR 박스
        qbox_h = 188*S
        qx0, qy0 = x, cur
        qx1, qy1 = card_w - pad, cur + qbox_h
        if not measure:
            d.rounded_rectangle([qx0, qy0, qx1, qy1], radius=10*S, outline=BORDER, width=S, fill=WHITE)
            qr = fake_qr(160*S)
            card.paste(qr, (int((qx0+qx1)/2 - 80*S), int(cur + 14*S)))
        cur = qy1 + 18*S
        # 수동 키 안내
        if not measure:
            d.text((x, cur), "QR을 스캔할 수 없으면 이 키를 수동 입력하세요:",
                   font=font(10*S), fill=MUTED)
        cur += 20*S
        # 시크릿 키 박스 (점선)
        kbox = [x, cur, card_w - pad, cur + 34*S]
        if not measure:
            d.rounded_rectangle(kbox, radius=8*S, fill=(0xF8, 0xF9, 0xFA))
            dashed_rect(d, kbox, (0xAD, 0xB5, 0xBD), width=S, dash=7*S, gap=4*S)
            d.text((x + 12*S, cur + 8*S), "JBSW Y3DP EHPK 3PXP 7QWE 2RTU",
                   font=font(12*S, mono=True), fill=DARKT)
        cur += 34*S + 26*S
        # ② 안내
        if not measure:
            d.text((x, cur), "② 앱에 표시된 6자리 코드를 입력해 등록을 완료하세요",
                   font=font(12*S, bold=True), fill=DARKT)
        cur += 28*S
        # OTP 입력박스
        ibox = [x, cur, card_w - pad, cur + 44*S]
        if not measure:
            d.rounded_rectangle(ibox, radius=8*S, outline=(0xCE, 0xD4, 0xDA), width=S, fill=WHITE)
            text_center(d, (x + card_w - pad)/2, cur + 8*S, "·  ·  ·  ·  ·  ·",
                        font(20*S, bold=True), (0xAD, 0xB5, 0xBD))
        cur += 44*S + 18*S
        # 버튼
        btn = [x, cur, card_w - pad, cur + 46*S]
        if not measure:
            grad = diag_gradient(int(btn[2]-btn[0]), 46*S, NAVY1, NAVY2)
            bm = rounded_mask(grad.width, grad.height, 8*S)
            card.paste(grad, (int(btn[0]), int(btn[1])), bm)
            text_center(d, (btn[0]+btn[2])/2, cur + 12*S, "✓  등록 완료", font(13*S, bold=True), WHITE)
        cur += 46*S + 28*S
        return cur - (0 if measure else (header_h + 22*S)) + (22*S if measure else 0)
    img = make_card(card_w, body, header, header_h)
    p = os.path.join(IMGDIR, "screen2_setup.png")
    img.save(p)
    return p

# ---------- 화면 ⑥ : 로그인 2단계 인증 ----------
def render_login():
    card_w = 410 * S
    header_h = 150 * S
    def header(card):
        d = ImageDraw.Draw(card)
        shield_lock(card, card_w/2, 52*S, 40*S)
        text_center(d, card_w/2, 86*S, "2단계 인증", font(20*S, bold=True), WHITE)
        # 부제 (2줄 wrap)
        sub = font(11*S)
        line1 = "인증 앱(Google Authenticator 등)에 표시된"
        line2 = "6자리 코드를 입력하세요."
        text_center(d, card_w/2, 116*S, line1, sub, (230, 236, 245))
        text_center(d, card_w/2, 132*S, line2, sub, (230, 236, 245))
        lang_pill(card, card_w - 12*S, 12*S)
    def body(card, measure=False):
        pad = 30 * S
        x = pad
        d = ImageDraw.Draw(card)
        cur = (24*S) if measure else (header_h + 24*S)
        base = 0 if measure else header_h
        # 라벨
        if not measure:
            d.text((x, cur), "인증 코드", font=font(12*S, bold=True), fill=DARKT)
        cur += 24*S
        # 입력박스
        ibox = [x, cur, card_w - pad, cur + 46*S]
        if not measure:
            d.rounded_rectangle(ibox, radius=8*S, outline=(0xCE, 0xD4, 0xDA), width=S, fill=WHITE)
            text_center(d, (x + card_w - pad)/2, cur + 9*S, "·  ·  ·  ·  ·  ·",
                        font(20*S, bold=True), (0xAD, 0xB5, 0xBD))
        cur += 46*S + 26*S
        # 버튼
        btn = [x, cur, card_w - pad, cur + 48*S]
        if not measure:
            grad = diag_gradient(int(btn[2]-btn[0]), 48*S, NAVY1, NAVY2)
            bm = rounded_mask(grad.width, grad.height, 8*S)
            card.paste(grad, (int(btn[0]), int(btn[1])), bm)
            text_center(d, (btn[0]+btn[2])/2, cur + 13*S, "✓  확인", font(13*S, bold=True), WHITE)
        cur += 48*S + 20*S
        # 돌아가기 링크
        if not measure:
            text_center(d, card_w/2, cur, "←  로그인으로 돌아가기", font(10.5*S), MUTED)
        cur += 24*S
        return cur - base + (0 if not measure else 0)
    img = make_card(card_w, body, header, header_h)
    p = os.path.join(IMGDIR, "screen6_login.png")
    img.save(p)
    return p

IMG_SETUP = render_setup()
IMG_LOGIN = render_login()
print("rendered:", IMG_SETUP, IMG_LOGIN)

# ===================================================================
#  PART 2.  워드 문서 생성
# ===================================================================
NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2D, 0x5A, 0xA0)
GRAY = RGBColor(0x6C, 0x75, 0x7D)
DARK = RGBColor(0x21, 0x25, 0x29)
KFONT = "맑은 고딕"

doc = Document()

def set_font(run, size=11, bold=False, color=DARK, font_name=KFONT):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.color.rgb = color; run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
        rfonts.set(qn(a), font_name)

normal = doc.styles['Normal']
normal.font.name = KFONT; normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), KFONT)

sec = doc.sections[0]
sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.3); sec.right_margin = Cm(2.3)

def shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hex_color)
    tcpr.append(sh)

def set_borders(cell, color="CCCCCC", sz="4", style="single"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), style); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tcpr.append(borders)

def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), size=15, bold=True, color=NAVY)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '2D5AA0')
    pbdr.append(bottom); ppr.append(pbdr)
    return p

def para(runs, align=None, space_after=4, space_before=0, indent=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    for text, kw in runs:
        set_font(p.add_run(text), **kw)
    return p

def screenshot_box(label, desc):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0); cell.width = Cm(15.5)
    shade(cell, "F4F7FB"); set_borders(cell, color="9DB8DA", sz="6", style="dashed")
    p0 = cell.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(14); p0.paragraph_format.space_after = Pt(2)
    set_font(p0.add_run(label), size=11, bold=True, color=NAVY)
    p1 = cell.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(14)
    set_font(p1.add_run(desc), size=9.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def image_box(img_path, caption, width_cm):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(img_path, width=Cm(width_cm))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    set_font(c.add_run(caption), size=9, color=GRAY)

def note_box(lines, fill="FFF8E1", border="E0C200"):
    tbl = doc.add_table(rows=1, cols=1); cell = tbl.cell(0, 0); cell.width = Cm(15.5)
    shade(cell, fill); set_borders(cell, color=border, sz="4", style="single")
    for i, (text, kw) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4 if i == 0 else 1)
        p.paragraph_format.space_after = Pt(4 if i == len(lines)-1 else 1)
        set_font(p.add_run(text), **kw)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---- 제목 ----
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
set_font(title.add_run("Google OTP(2단계 인증) 등록 안내"), size=22, bold=True, color=NAVY)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(10)
set_font(sub.add_run("연결 재무보고 통합 시스템"), size=12, color=BLUE)

para([("연결 재무보고 통합 시스템은 보안을 위해 ", {}),
      ("2단계 인증(OTP)", {"bold": True, "color": NAVY}),
      ("을 사용합니다. 아이디·비밀번호 외에, 휴대폰 앱에 표시되는 ", {}),
      ("6자리 숫자", {"bold": True, "color": NAVY}),
      ("를 한 번 더 입력해야 로그인됩니다.", {})], space_after=4)
para([("최초 1회만 등록", {"bold": True, "color": NAVY}),
      ("하면 되고, 이후에는 로그인할 때 코드만 입력하면 됩니다.", {})], space_after=6)
note_box([("소요 시간: 약 2~3분   |   준비물: 스마트폰",
           {"size": 10.5, "bold": True, "color": NAVY})], fill="EAF1FB", border="9DB8DA")

# ---- 1단계 ----
heading("1단계.  휴대폰에 인증 앱 설치하기")
para([("스마트폰에서 아래 앱 중 하나를 설치하세요. ", {}),
      ("(Google Authenticator 권장)", {"bold": True, "color": NAVY})], space_after=6)
t = doc.add_table(rows=3, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [Cm(4.5), Cm(11.0)]
rows_data = [("휴대폰 종류", "설치 방법", True),
             ("아이폰(iPhone)", "App Store 에서  「Google Authenticator」  검색 후 설치", False),
             ("안드로이드", "Play 스토어 에서  「Google Authenticator」  검색 후 설치", False)]
for ri, (c1, c2, head) in enumerate(rows_data):
    for ci, txt in enumerate((c1, c2)):
        cell = t.cell(ri, ci); cell.width = widths[ci]
        set_borders(cell, color="C9D6E5", sz="4")
        if head: shade(cell, "1F3864")
        elif ri % 2 == 0: shade(cell, "F4F7FB")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        if ci == 0: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(txt), size=10.5, bold=head or ci == 0,
                 color=RGBColor(0xFF, 0xFF, 0xFF) if head else DARK)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
note_box([("Microsoft Authenticator, Authy 앱도 똑같이 사용할 수 있습니다.",
           {"size": 10, "color": GRAY})], fill="F4F7FB", border="C9D6E5")
screenshot_box("📷 캡쳐 ①", "앱스토어에서 Google Authenticator 를 검색한 화면")

# ---- 2단계 ----
heading("2단계.  시스템에서 등록 화면 열기")
para([("PC에서 시스템에 ", {}), ("아이디·비밀번호로 로그인", {"bold": True, "color": NAVY}),
      ("하면, 아직 OTP를 등록하지 않은 경우 아래와 같은 ", {}),
      ("「2단계 인증(OTP) 등록」", {"bold": True, "color": NAVY}),
      (" 화면이 자동으로 나타납니다.", {})], space_after=6)
note_box([("“보안 정책에 따라 OTP 등록 후 시스템을 이용할 수 있습니다” 라는 안내가 보이면 정상입니다.",
           {"size": 10, "color": DARK}),
          ("등록을 마쳐야 다음 화면으로 넘어갑니다.", {"size": 10, "color": GRAY})])
image_box(IMG_SETUP, "▲ 시스템의 「2단계 인증(OTP) 등록」 화면 (실제 화면 예시)", 11.5)

# ---- 3단계 ----
heading("3단계.  앱으로 QR 코드 스캔하기")
for a, b, c in [("①  휴대폰에서 ", "Google Authenticator 앱", " 을 엽니다."),
                ("②  화면 아래(또는 오른쪽 아래)의 ", "＋ (코드 추가 / QR 코드 스캔)", " 버튼을 누릅니다."),
                ("③  ", "「QR 코드 스캔」", " 을 선택합니다."),
                ("④  휴대폰 카메라로 PC 화면의 ", "QR 코드", " 를 비춥니다.")]:
    para([(a, {}), (b, {"bold": True, "color": NAVY}), (c, {})], space_after=3, indent=0.3)
para([("→ 스캔되면 앱에 ", {}), ("“연결재무보고시스템”", {"bold": True, "color": NAVY}),
      (" 항목이 생기고, ", {}), ("6자리 숫자", {"bold": True, "color": NAVY}),
      ("가 표시됩니다.", {})], space_before=4, space_after=6, indent=0.3)
screenshot_box("📷 캡쳐 ③", "앱에서 ＋ 버튼 → 「QR 코드 스캔」 을 누르는 화면")
screenshot_box("📷 캡쳐 ④", "스캔 후 6자리 코드가 표시된 앱 화면")
note_box([("QR 코드가 스캔되지 않을 때", {"size": 10.5, "bold": True, "color": RGBColor(0x8A, 0x6D, 0x00)}),
          ("등록 화면의 QR 코드 아래에 있는 키(영문·숫자 문자열)를 앱에서 직접 입력해도 됩니다.",
           {"size": 10, "color": DARK}),
          ("앱에서  ＋ → 「설정 키 입력」  을 선택한 뒤, 계정 이름은 아무거나, 키 칸에 그 문자열을 입력하세요.",
           {"size": 10, "color": DARK})])

# ---- 4단계 ----
heading("4단계.  6자리 코드 입력해서 등록 완료")
for a, b, c in [("①  앱에 표시된 ", "6자리 숫자", " 를 확인합니다."),
                ("②  PC 등록 화면의 입력칸에 그 ", "6자리 숫자", " 를 입력합니다."),
                ("③  ", "[등록 완료]", " 버튼을 누릅니다.")]:
    para([(a, {}), (b, {"bold": True, "color": NAVY}), (c, {})], space_after=3, indent=0.3)
para([("✓ ", {"bold": True, "color": RGBColor(0x1E, 0x7E, 0x34)}),
      ("“등록 완료”되면 메인 화면으로 들어갑니다. 이제 등록이 끝났습니다!",
       {"bold": True, "color": RGBColor(0x1E, 0x7E, 0x34)})], space_before=4, space_after=6, indent=0.3)
screenshot_box("📷 캡쳐 ⑤", "6자리 코드를 입력하고 [등록 완료] 를 누르는 화면")
note_box([("코드는 30초마다 자동으로 바뀝니다.", {"size": 10.5, "bold": True, "color": RGBColor(0x8A, 0x6D, 0x00)}),
          ("입력 중 숫자가 바뀌면, 앱에 보이는 최신 6자리로 다시 입력하세요.", {"size": 10, "color": DARK})])

# ---- 다음 로그인 ----
heading("다음 로그인부터는?")
para([("앞으로 로그인할 때마다 아이디·비밀번호를 입력한 뒤, ", {}),
      ("「2단계 인증」", {"bold": True, "color": NAVY}),
      (" 화면이 나오면 앱을 열어 그때그때 표시되는 ", {}),
      ("6자리 코드", {"bold": True, "color": NAVY}),
      ("를 입력하면 됩니다. ", {}), ("(재등록 불필요)", {"color": GRAY})], space_after=6)
image_box(IMG_LOGIN, "▲ 로그인 시 나타나는 「2단계 인증」 코드 입력 화면 (실제 화면 예시)", 10.0)

# ---- FAQ ----
heading("자주 묻는 질문")
for q, a in [
    ("Q.  휴대폰을 바꿨어요 / 앱을 지웠어요 / 코드가 계속 틀려요",
     "관리자에게 “2단계 인증(2FA) 초기화”를 요청하세요. 초기화 후 1~4단계를 다시 등록하면 됩니다."),
    ("Q.  코드를 입력해도 “올바르지 않습니다”라고 나와요",
     "휴대폰 시간이 자동(네트워크 시간)으로 맞춰져 있는지 확인하세요. 시간이 어긋나면 코드가 틀릴 수 있습니다. "
     "그리고 현재 앱에 떠 있는 최신 6자리인지 다시 확인하세요."),
    ("Q.  로그인 코드 입력 화면이 사라졌어요",
     "코드 입력은 5분 이내에 마쳐야 합니다. 시간이 지나면 로그인부터 다시 하세요.")]:
    para([(q, {"bold": True, "color": NAVY, "size": 10.5})], space_before=5, space_after=2)
    para([("    " + a, {"size": 10.5, "color": DARK})], space_after=3)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(foot.add_run("문의: 시스템 관리자"), size=9, color=GRAY)

out = r"C:\패키지프로그램\Google_OTP_등록안내.docx"
doc.save(out)
print("saved:", out)
