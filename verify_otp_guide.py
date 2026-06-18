# -*- coding: utf-8 -*-
import docx
d = docx.Document(r"C:\패키지프로그램\Google_OTP_등록안내.docx")
print("paragraphs:", len(d.paragraphs), "tables:", len(d.tables))
print("title:", d.paragraphs[0].runs[0].text)
caps = []
for t in d.tables:
    runs = t.cell(0, 0).paragraphs[0].runs
    if runs and "캡쳐" in runs[0].text:
        caps.append(runs[0].text)
print("screenshot boxes:", caps)
